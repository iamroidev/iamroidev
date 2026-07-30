#!/usr/bin/env python3
"""
Rebuild CV DOCX from the student CV template with proper paragraph styles.
Starts from the template to inherit all styles, fonts, and numbering definitions.

Template structure uses:
- Heading 2 style for institution/title names (bold, Palatino Linotype)
- List Paragraph style + numPr XML for bulleted achievements (Tahoma font)
- Each sub-line as a separate paragraph within the table cell
"""

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "STUDENT CV TEMPLATE.docx"
OUT = ROOT / "cv" / "CE3_FCM4100819823_OPOKU-Richard-Kwaku_CV.docx"

# numId=1 in the template (abstractNumId=2) is the bullet list used in example content
BULLET_NUM_ID = "1"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _tahoma_rpr(bold=False):
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Tahoma")
    rFonts.set(qn("w:hAnsi"), "Tahoma")
    rFonts.set(qn("w:cs"), "Tahoma")
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    return rPr


def _make_run(text, bold=False):
    r = OxmlElement("w:r")
    r.append(_tahoma_rpr(bold))
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def make_para(text, style_id="Normal", bold=False, bullet=False):
    """Build a <w:p> element with the given style and text."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)

    if bullet:
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), BULLET_NUM_ID)
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "40")
        spacing.set(qn("w:line"), "260")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

    p.append(pPr)
    p.append(_make_run(text, bold=bold))
    return p


def make_heading2_para(text):
    """Institution/title name: Heading 2 style (bold, Palatino Linotype)."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Heading2")
    pPr.append(pStyle)
    p.append(pPr)
    # Heading 2 already defines Palatino Linotype bold via the style; just add a run
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


# ---------------------------------------------------------------------------
# Cell helpers
# ---------------------------------------------------------------------------

def clear_cell(cell):
    """Remove every paragraph from a table cell."""
    tc = cell._tc
    for p in list(tc.findall(qn("w:p"))):
        tc.remove(p)


def fill_date_cell(cell, date_text):
    clear_cell(cell)
    cell._tc.append(make_para(date_text, "Normal"))


def fill_entry_cell(cell, title, organisation, bullets=None):
    """
    Right-hand cell structure:
      title        -> Heading 2 (bold institution/title name)
      organisation -> Normal (subtitle/org line)
      bullets      -> List Paragraph with numPr, one per bullet
    """
    clear_cell(cell)
    tc = cell._tc
    tc.append(make_heading2_para(title))
    if organisation:
        tc.append(make_para(organisation, "Normal"))
    if bullets:
        for b in bullets:
            tc.append(make_para(b, "ListParagraph", bullet=True))


def fill_skills_cell(cell, text):
    """Skills table right-hand cell: plain Normal text."""
    clear_cell(cell)
    cell._tc.append(make_para(text, "Normal"))


def fill_skills_label_cell(cell, text):
    """Skills table left-hand label cell: bold Normal."""
    clear_cell(cell)
    cell._tc.append(make_para(text, "Normal", bold=True))


# ---------------------------------------------------------------------------
# Row helpers
# ---------------------------------------------------------------------------

def set_row(row, date_text, title, organisation, bullets=None):
    fill_date_cell(row.cells[0], date_text)
    fill_entry_cell(row.cells[1], title, organisation, bullets)


def delete_all_rows(table):
    tbl = table._tbl
    for tr in list(tbl.findall(qn("w:tr"))):
        tbl.remove(tr)


def add_entry_row(table, date_text, title, organisation, bullets=None):
    """Append a new entry row, cloning the table's first row for formatting."""
    if table.rows:
        new_tr = deepcopy(table.rows[0]._tr)
    else:
        new_tr = OxmlElement("w:tr")
        for _ in range(2):
            tc = OxmlElement("w:tc")
            new_tr.append(tc)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    set_row(row, date_text, title, organisation, bullets)


# ---------------------------------------------------------------------------
# Body text paragraph helpers (for PROFESSIONAL SUMMARY outside tables)
# ---------------------------------------------------------------------------

def replace_body_text_paragraph(para, text):
    """Replace the text in an existing Body Text paragraph."""
    para.clear()
    r = para.add_run(text)
    r.font.name = "Tahoma"


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def build():
    shutil.copy(TEMPLATE, OUT)
    doc = Document(str(OUT))

    # --- Page header text box: "Curriculum Vitae of Your Name" ---
    from docx.oxml.ns import qn as _qn
    for rel_id, rel in doc.part.rels.items():
        if "header" in str(rel.reltype).lower():
            hdr = rel.target_part
            for t in hdr._element.findall(".//" + _qn("w:t")):
                if t.text == "Your Name":
                    t.text = "Richard Kwaku Opoku"

    # --- Body header: update name ---
    # Paragraph [0] is Title style (FULL NAME)
    title_para = doc.paragraphs[0]
    for run in title_para.runs:
        run.text = ""
    if title_para.runs:
        title_para.runs[0].text = "RICHARD KWAKU OPOKU"
    else:
        title_para.add_run("RICHARD KWAKU OPOKU")

    # Paragraph [1]: contact line 1
    p1 = doc.paragraphs[1]
    for run in p1.runs:
        run.text = ""
    if p1.runs:
        p1.runs[0].text = (
            "Mobile Number: +233 55 150 0736  |  Email: richardkwakuopoku06@gmail.com"
        )
    else:
        p1.add_run("Mobile Number: +233 55 150 0736  |  Email: richardkwakuopoku06@gmail.com")

    # Paragraph [2]: address
    p2 = doc.paragraphs[2]
    for run in p2.runs:
        run.text = ""
    if p2.runs:
        p2.runs[0].text = "Address: UMaT, Tarkwa, Ghana"
    else:
        p2.add_run("Address: UMaT, Tarkwa, Ghana")

    # Paragraph [3]: LinkedIn
    p3 = doc.paragraphs[3]
    for run in p3.runs:
        run.text = ""
    if p3.runs:
        p3.runs[0].text = (
            "LinkedIn: linkedin.com/in/richardkwakuopoku982"
            "  |  GitHub: github.com/iamroidev"
            "  |  Site: richardkwakuopoku.site"
        )
    else:
        p3.add_run(
            "LinkedIn: linkedin.com/in/richardkwakuopoku982"
            "  |  GitHub: github.com/iamroidev"
            "  |  Site: richardkwakuopoku.site"
        )

    # --- Professional Summary ---
    # Find and replace the Body Text paragraph after PROFESSIONAL SUMMARY heading
    summary_text = (
        "Computer Science and Engineering student at UMaT (Year 3, 80.87 CWA, First Class standing). "
        "AWS re/Start graduate with strong cloud software engineering foundations, hands-on security "
        "training, and experience building and deploying live web applications "
        "(React, Next.js, Node.js, Python, and AWS EC2/Amplify)."
    )
    in_summary = False
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and "PROFESSIONAL SUMMARY" in para.text.upper():
            in_summary = True
            continue
        if in_summary and para.style.name == "Heading 1":
            break
        if in_summary and para.style.name in ("Body Text", "Normal") and para.text.strip():
            # Replace first non-empty body text paragraph with our content
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = summary_text
            else:
                para.add_run(summary_text)
            # Clear remaining instruction paragraphs
            in_summary = False  # only need the first

    # Also clear remaining instruction Body Text paragraphs in summary section
    in_summary = False
    instruction_cleared = False
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and "PROFESSIONAL SUMMARY" in para.text.upper():
            in_summary = True
            instruction_cleared = False
            continue
        if in_summary and para.style.name == "Heading 1":
            break
        if in_summary and para.style.name in ("Body Text", "Normal"):
            if not instruction_cleared and para.text.strip():
                instruction_cleared = True
                continue  # skip the one we already set above
            # Clear all other instruction paragraphs
            for run in para.runs:
                run.text = ""

    # --- Education (Table 0) ---
    edu_table = doc.tables[0]
    delete_all_rows(edu_table)
    add_entry_row(
        edu_table,
        "2024 – Present",
        "UNIVERSITY OF MINES AND TECHNOLOGY – TARKWA, GHANA",
        "BSc Computer Science and Engineering  |  Expected 2027",
        bullets=[
            "CWA: 80.87 (First Class standing)",
            "Relevant coursework: Data Structures & Algorithms, Software Engineering, Artificial Intelligence, "
            "Advanced Database Systems, Operating Systems, Embedded System Design, Web Programming, "
            "Object-Oriented Programming",
        ],
    )
    add_entry_row(
        edu_table,
        "2019 – 2022",
        "KUMASI ACADEMY – KUMASI, GHANA",
        "West Africa Senior Secondary School Certificate (WASSCE)",
        bullets=["Programme: General Science"],
    )

    # --- Work Experience (Table 1) ---
    exp_table = doc.tables[1]
    delete_all_rows(exp_table)
    add_entry_row(
        exp_table,
        "Jan. 2026 – Apr. 2026",
        "AWS RE/START CLOUD PRACTITIONER GRADUATE",
        "AmaliTech – Online",
        bullets=[
            "Completed a 12-week intensive training programme covering cloud architecture, IAM policies, "
            "VPC configuration, EC2, S3 bucket security, Linux systems administration, Python scripting, "
            "networking, and relational/NoSQL databases.",
        ],
    )

    # --- Projects and Leadership (Table 2) ---
    proj_table = doc.tables[2]
    delete_all_rows(proj_table)
    add_entry_row(
        proj_table,
        "2026",
        "VOTEEQ – PAID NOMINEE VOTING & EVENT TICKETING PLATFORM",
        "Full-Stack Developer & Architect  |  voteeq.online",
        bullets=[
            "Built and deployed a paid nominee voting and event ticketing platform serving 45 award "
            "categories, 63 nominees, and 7,891 votes (GHS 7,891 revenue) for the official ACSES "
            "Praemia Pro Virtute Dinner & Awards Night.",
            "Implemented Paystack hosted checkout, real-time leaderboard updates via WebSockets, "
            "libSQL database, nominee dashboards, and admin tooling; frontend on Vercel, API on AWS EC2.",
        ],
    )
    add_entry_row(
        proj_table,
        "2025",
        "SCHOLAR – AI SCHOLARSHIP MATCHING ENGINE",
        "Full-Stack Developer  |  schorla.vercel.app",
        bullets=[
            "Built an NLP-powered matching engine that parses CVs, transcripts, and SOPs to align "
            "student profiles with verified funding opportunities.",
            "Developed with Next.js, LlamaCloud, Supabase, and Stripe billing.",
        ],
    )
    add_entry_row(
        proj_table,
        "2025",
        "INSIGHTFLOW – AI-POWERED READING ASSISTANT",
        "Full-Stack Developer  |  appinsightflow.vercel.app",
        bullets=[
            "Built a RAG pipeline that converts books and documents into structured learning paths, "
            "generating guided text summaries and synthesized audio streams.",
            "Developed with React, Vite, TypeScript, and Supabase.",
        ],
    )
    add_entry_row(
        proj_table,
        "2026",
        "KAIROVA – PREMIUM STREETWEAR E-COMMERCE SHOWCASE",
        "Full-Stack Developer  |  kai-rova.vercel.app",
        bullets=[
            "Built a premium clothing catalog and brand experience showcase for a streetwear brand.",
            "Implemented category catalog filters, interactive detail galleries, size selector drawers, "
            "and dynamic user query forms. Built with Next.js 15, TypeScript, Tailwind CSS, Framer Motion.",
        ],
    )
    add_entry_row(
        proj_table,
        "2026",
        "AWS RE/START GRADUATE HACKATHON – AI MEDICAL QUEUE APPLICATION",
        "Team Hackathon Project  |  AmaliTech",
        bullets=[
            "Collaborated in a team to design and build a hospital queue management system.",
            "Implemented user authentication via AWS Cognito, serverless endpoints with Lambda + "
            "API Gateway, DynamoDB database, and Amplify frontend hosting.",
        ],
    )
    add_entry_row(
        proj_table,
        "2026",
        "AMALITECH CODING CLUB ORGANIZER",
        "Coding Club  |  UMaT",
        bullets=[
            "Organized coding workshops, technical bootcamps, and peer-to-peer programming practice sessions.",
            "Coordinated student mentoring, algorithmic tutorials, and collaborative hacking events.",
        ],
    )
    add_entry_row(
        proj_table,
        "2026 – Present",
        "UMaT CYBERSECURITY CLUB – MEMBER",
        "Cybersecurity Club  |  UMaT",
        bullets=[
            "Participated in hands-on technical workshops focused on network scanning and enumeration "
            "using Nmap, and database vulnerability exploitation via SQL injection.",
        ],
    )
    add_entry_row(
        proj_table,
        "2024 – Present",
        "ROBOTICS CLUB (AAENICS) – IOT & EMBEDDED SYSTEMS",
        "UMaT",
        bullets=[
            "Completed hands-on training and built embedded prototypes using ESP32, MQTT protocols, "
            "Embedded C, and Raspberry Pi.",
        ],
    )

    # --- Workshops / Seminars (Table 3) ---
    ws_table = doc.tables[3]
    delete_all_rows(ws_table)
    add_entry_row(
        ws_table,
        "24th July, 2026",
        "PARTICIPANT, AMALITECH WORKSHOP – HOW TO ACTUALLY GET HIRED IN TECH",
        "",
    )

    # --- Relevant Skills (Table 4) ---
    skills_table = doc.tables[4]
    skills_data = [
        (
            "Technical / IT",
            "AWS (Amplify, S3, IAM, Lambda, API Gateway, EC2, VPC), Python, Node.js, "
            "Next.js, React, TypeScript, Tailwind CSS, PostgreSQL, Supabase, Prisma, "
            "Docker, Firebase, Vercel, GCP, REST APIs, Three.js",
        ),
        (
            "Soft skills",
            "Team Coordination, Public Speaking, Negotiation, Customer Service",
        ),
        (
            "Certifications",
            "AWS Certified Cloud Practitioner (CCP) – Amazon Web Services (2026)  |  "
            "Certified in Cybersecurity (CC) – ISC2 (2026)  |  "
            "Machine Learning Specialization – DeepLearning.AI & Stanford University (Coursera)",
        ),
        (
            "Languages",
            "English (Fluent), Twi (Fluent)",
        ),
        (
            "Hobbies and interests",
            "Building software products, researching cloud technologies and cybersecurity, "
            "reading about artificial intelligence, IoT and embedded systems, entrepreneurship",
        ),
    ]
    for i, (label, value) in enumerate(skills_data):
        if i < len(skills_table.rows):
            row = skills_table.rows[i]
            fill_skills_label_cell(row.cells[0], label)
            fill_skills_cell(row.cells[1], value)

    # --- Referees (Table 5) ---
    ref_table = doc.tables[5]
    delete_all_rows(ref_table)
    # Add a single row for "Available on request"
    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    # colspan spanning both columns
    tcPr = OxmlElement("w:tcPr")
    gridSpan = OxmlElement("w:gridSpan")
    gridSpan.set(qn("w:val"), "2")
    tcPr.append(gridSpan)
    tc.append(tcPr)
    tc.append(make_para("Available on request.", "Normal"))
    tr.append(tc)
    ref_table._tbl.append(tr)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
